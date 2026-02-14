"""
M11 table renderers — Synopsis table and Schedule of Activities (SoA) table.

Generates structured DOCX tables from USDM entities for:
  - §1.1.2 Synopsis (Overall Design) table
  - §1.3 Schedule of Activities table
"""

import json
import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _build_soa_data(usdm: Dict) -> Optional[Dict]:
    """Extract SoA table data from USDM for rendering.

    Returns a dict with:
      epochs: [{id, name}]
      encounters_by_epoch: {epoch_id: [{id, name, timing}]}
      groups: [{name, activities: [{id, name}]}]
      ungrouped: [{id, name}]
      cells: {(activity_id, encounter_id): mark}
      footnote_refs: {(activity_id, encounter_id): [letter, ...]}
      footnotes: [str, ...]
    """
    study = usdm.get('study', {})
    versions = study.get('versions', [{}])
    version = versions[0] if versions else {}
    design = (version.get('studyDesigns', [{}]) or [{}])[0]

    timelines = design.get('scheduleTimelines', [])
    if not timelines:
        return None

    timeline = timelines[0]
    instances = timeline.get('instances', [])
    if not instances:
        return None

    # Epochs
    epochs = design.get('epochs', [])
    epoch_list = [{'id': e.get('id', ''), 'name': e.get('name', '')}
                  for e in epochs if isinstance(e, dict)]

    # Encounters
    encounters = design.get('encounters', [])
    enc_map = {e.get('id', ''): e for e in encounters if isinstance(e, dict)}

    # Build encounters_by_epoch
    encounters_by_epoch: Dict[str, list] = {}
    for enc in encounters:
        if not isinstance(enc, dict):
            continue
        epoch_id = enc.get('epochId', '')
        if not epoch_id:
            # Try to find epoch from instances
            for inst in instances:
                if inst.get('encounterId') == enc.get('id'):
                    epoch_id = inst.get('epochId', '')
                    break
        if epoch_id:
            encounters_by_epoch.setdefault(epoch_id, []).append({
                'id': enc.get('id', ''),
                'name': enc.get('name', ''),
                'timing': enc.get('scheduledAtTimingId', ''),
            })

    # Activities
    activities = design.get('activities', [])
    act_map = {a.get('id', ''): a for a in activities if isinstance(a, dict)}

    # Activity groups
    groups_raw = design.get('activityGroups', [])
    grouped_ids = set()
    groups = []
    for g in groups_raw:
        if not isinstance(g, dict):
            continue
        child_ids = g.get('childIds', [])
        group_acts = []
        for cid in child_ids:
            act = act_map.get(cid)
            if act:
                group_acts.append({'id': cid, 'name': act.get('name', '')})
                grouped_ids.add(cid)
        if group_acts:
            groups.append({'name': g.get('name', ''), 'activities': group_acts})

    # Ungrouped activities
    ungrouped = []
    for act in activities:
        if isinstance(act, dict) and act.get('id') not in grouped_ids:
            # Only include activities that have ticks in the SoA
            ungrouped.append({'id': act.get('id', ''), 'name': act.get('name', '')})

    # Build cell marks from schedule instances
    cells: Dict[tuple, str] = {}
    for inst in instances:
        if not isinstance(inst, dict):
            continue
        enc_id = inst.get('encounterId', '')
        act_ids = inst.get('activityIds', [])
        for aid in act_ids:
            cells[(aid, enc_id)] = 'X'

    # Footnotes from extension attributes
    footnotes = []
    footnote_refs: Dict[tuple, list] = {}
    try:
        for ext in design.get('extensionAttributes', []):
            if isinstance(ext, dict) and ext.get('url', '').endswith('soaFootnotes'):
                raw = json.loads(ext.get('valueString', '[]'))
                if isinstance(raw, list):
                    for fn in raw:
                        if isinstance(fn, dict):
                            text = fn.get('text', fn.get('footnote', ''))
                            if text:
                                footnotes.append(text)
                            # Map footnote references to cells
                            ref_acts = fn.get('activityIds', [])
                            ref_encs = fn.get('encounterIds', [])
                            letter = chr(96 + len(footnotes))  # a, b, c, ...
                            for aid in ref_acts:
                                for eid in ref_encs:
                                    footnote_refs.setdefault((aid, eid), []).append(letter)
    except (json.JSONDecodeError, TypeError):
        pass

    return {
        'epochs': epoch_list,
        'encounters_by_epoch': encounters_by_epoch,
        'groups': groups,
        'ungrouped': ungrouped,
        'cells': cells,
        'footnote_refs': footnote_refs,
        'footnotes': footnotes,
    }


# ── Cell styling helpers ─────────────────────────────────────────────

_HEADER_SHADE = 'D9E2F3'   # Light blue-grey for epoch/encounter headers
_GROUP_SHADE  = 'E2EFDA'   # Light green for group separator rows
_FONT_NAME    = 'Arial Narrow'


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): hex_color,
    })
    tc_pr.append(shading)


def _set_cell_vertical_alignment(cell, align: str = 'center'):
    """Set vertical alignment of a table cell ('top', 'center', 'bottom')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.makeelement(qn('w:vAlign'), {qn('w:val'): align})
    tc_pr.append(v_align)


def _set_cell_margins(cell, top=30, bottom=30, left=40, right=40):
    """Set cell margins in twips (1/20 of a point)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.makeelement(qn('w:tcMar'), {})
    for side, val in [('top', top), ('bottom', bottom),
                      ('start', left), ('end', right)]:
        m = margins.makeelement(qn(f'w:{side}'), {
            qn('w:w'): str(val), qn('w:type'): 'dxa',
        })
        margins.append(m)
    tc_pr.append(margins)


def _set_col_width(cell, width_emu: int):
    """Set explicit column width on a cell via XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.makeelement(qn('w:tcW'), {
        qn('w:w'): str(int(width_emu / 635)),  # EMU → twips
        qn('w:type'): 'dxa',
    })
    tc_pr.append(tc_w)


def _mark_row_as_header(row):
    """Mark a table row to repeat on every page (DOCX tblHeader)."""
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.makeelement(qn('w:tblHeader'), {})
    tr_pr.append(header)


def _set_row_height(row, height_pt: float):
    """Set exact row height."""
    tr_pr = row._tr.get_or_add_trPr()
    rh = tr_pr.makeelement(qn('w:trHeight'), {
        qn('w:val'): str(int(height_pt * 20)),  # pt → twips
        qn('w:hRule'): 'atLeast',
    })
    tr_pr.append(rh)


def _style_header_cell(cell, text: str, font_size: float = 8,
                       bold: bool = True, center: bool = True):
    """Write text into a header cell with standard styling."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = _FONT_NAME
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _shade_cell(cell, _HEADER_SHADE)
    _set_cell_vertical_alignment(cell)
    _set_cell_margins(cell)


# ── Main SoA table builder ───────────────────────────────────────────

def _add_soa_table(doc: Document, usdm: Dict) -> bool:
    """Add Schedule of Activities table to the document.

    Creates a landscape-oriented section with a grid table showing
    activities (rows) × encounters (columns) with X marks.
    Features:
      - Merged epoch header row with shading
      - Encounter header row with shading
      - Group separator rows with green shading
      - Empty activities (no ticks) are filtered out
      - Header rows repeat on every page
      - Footnotes rendered below the table

    Returns True if table was added, False if no SoA data available.
    """
    soa = _build_soa_data(usdm)
    if not soa or not soa['cells']:
        return False

    epochs = soa['epochs']
    encounters_by_epoch = soa['encounters_by_epoch']
    groups = soa['groups']
    ungrouped = soa['ungrouped']
    cells = soa['cells']
    footnotes = soa['footnotes']
    footnote_refs = soa['footnote_refs']

    # Flatten encounters in epoch order
    ordered_encounters: List[Dict] = []
    for epoch in epochs:
        epoch_encs = encounters_by_epoch.get(epoch['id'], [])
        ordered_encounters.extend(epoch_encs)

    if not ordered_encounters:
        return False

    enc_ids = [e['id'] for e in ordered_encounters]

    # Filter activities: only keep those with at least one tick
    def _has_ticks(act_id: str) -> bool:
        return any(cells.get((act_id, eid)) for eid in enc_ids)

    # Build ordered row descriptors: ('group', group_name) or ('activity', act_dict)
    row_descriptors: List[tuple] = []
    for group in groups:
        active = [a for a in group['activities'] if _has_ticks(a['id'])]
        if active:
            row_descriptors.append(('group', group['name']))
            for a in active:
                row_descriptors.append(('activity', a))
    active_ungrouped = [a for a in ungrouped if _has_ticks(a['id'])]
    if active_ungrouped:
        row_descriptors.append(('group', 'Other Assessments'))
        for a in active_ungrouped:
            row_descriptors.append(('activity', a))

    if not row_descriptors:
        return False

    activity_count = sum(1 for t, _ in row_descriptors if t == 'activity')

    # ── Page layout: landscape ────────────────────────────────────
    new_section = doc.add_section(WD_ORIENT.LANDSCAPE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Inches(11)
    new_section.page_height = Inches(8.5)
    new_section.top_margin = Inches(0.5)
    new_section.bottom_margin = Inches(0.5)
    new_section.left_margin = Inches(0.5)
    new_section.right_margin = Inches(0.5)

    # ── Create table ──────────────────────────────────────────────
    # Rows: 2 header rows + row_descriptors
    n_header_rows = 2
    n_rows = n_header_rows + len(row_descriptors)
    n_cols = 1 + len(ordered_encounters)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    usable_width = Inches(10)  # 11" - 0.5" margins each side
    act_col_width = Inches(2.5)
    enc_col_width_emu = int((usable_width - act_col_width) / max(len(ordered_encounters), 1))

    # Apply widths to every row (DOCX requires per-cell width)
    for row in table.rows:
        _set_col_width(row.cells[0], int(act_col_width))
        for j in range(1, n_cols):
            _set_col_width(row.cells[j], enc_col_width_emu)

    # ── Row 0: Epoch headers (merged across encounters) ───────────
    _mark_row_as_header(table.rows[0])
    _set_row_height(table.rows[0], 16)
    _style_header_cell(table.cell(0, 0), 'Study Period', font_size=8)

    epoch_col = 1
    for epoch in epochs:
        epoch_encs = encounters_by_epoch.get(epoch['id'], [])
        if not epoch_encs:
            continue
        start_col = epoch_col
        end_col = epoch_col + len(epoch_encs) - 1

        cell = table.cell(0, start_col)
        if end_col > start_col:
            cell.merge(table.cell(0, end_col))
        _style_header_cell(cell, epoch['name'], font_size=8)

        epoch_col += len(epoch_encs)

    # ── Row 1: Encounter names ────────────────────────────────────
    _mark_row_as_header(table.rows[1])
    _set_row_height(table.rows[1], 14)
    _style_header_cell(table.cell(1, 0), 'Assessment', font_size=7.5,
                       center=False)

    for j, enc in enumerate(ordered_encounters):
        _style_header_cell(table.cell(1, j + 1), enc['name'], font_size=7)

    # ── Data rows ─────────────────────────────────────────────────
    for i, (row_type, payload) in enumerate(row_descriptors):
        row_idx = n_header_rows + i
        row = table.rows[row_idx]

        if row_type == 'group':
            # Group separator row — merge all columns, green shading
            cell = table.cell(row_idx, 0)
            cell.merge(table.cell(row_idx, n_cols - 1))
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(payload)
            run.bold = True
            run.font.size = Pt(7.5)
            run.font.name = _FONT_NAME
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _shade_cell(cell, _GROUP_SHADE)
            _set_cell_vertical_alignment(cell)
            _set_cell_margins(cell)
            _set_row_height(row, 14)
        else:
            # Activity data row
            act = payload
            # Activity name
            name_cell = table.cell(row_idx, 0)
            name_cell.text = ''
            p = name_cell.paragraphs[0]
            run = p.add_run(act['name'])
            run.font.size = Pt(7)
            run.font.name = _FONT_NAME
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _set_cell_vertical_alignment(name_cell)
            _set_cell_margins(name_cell)

            # Tick marks
            for j, enc in enumerate(ordered_encounters):
                cell = table.cell(row_idx, j + 1)
                mark = cells.get((act['id'], enc['id']), '')
                refs = footnote_refs.get((act['id'], enc['id']), [])

                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                if mark:
                    run = p.add_run(mark)
                    run.font.size = Pt(8)
                    run.font.name = _FONT_NAME
                if refs:
                    ref_text = ''.join(refs)
                    run = p.add_run(ref_text)
                    run.font.size = Pt(6)
                    run.font.name = _FONT_NAME
                    run.font.superscript = True

                _set_cell_vertical_alignment(cell)
                _set_cell_margins(cell, top=20, bottom=20, left=20, right=20)

            _set_row_height(row, 12)

    # ── Footnotes below table ─────────────────────────────────────
    if footnotes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        for idx, fn in enumerate(footnotes):
            letter = chr(97 + idx)  # a, b, c, ...
            run = p.add_run(f'{letter}. {fn}\n')
            run.font.size = Pt(7)
            run.font.name = _FONT_NAME
            run.italic = True

    # ── Switch back to portrait ───────────────────────────────────
    new_section = doc.add_section(WD_ORIENT.PORTRAIT)
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)

    logger.info(f"  ✓ SoA table: {activity_count} activities × "
                f"{len(ordered_encounters)} encounters, "
                f"{len(row_descriptors) - activity_count} group headers")
    return True


def _add_synopsis_table(doc: Document, usdm: Dict) -> bool:
    """Add Synopsis Overall Design table to the document.

    Creates a two-column table with field labels and values from
    the _compose_synopsis data.

    Returns True if table was added, False if no synopsis data available.
    """
    from .composers import _compose_synopsis

    synopsis_text = _compose_synopsis(usdm)
    if not synopsis_text.strip():
        return False

    # Parse the structured field list
    fields = []
    for line in synopsis_text.split('\n'):
        line = line.strip()
        if not line or line.startswith('**Overall Design**'):
            continue
        # Parse "  **Label**: Value" format
        if line.startswith('**') and '**:' in line:
            parts = line.split('**:', 1)
            label = parts[0].replace('**', '').strip()
            value = parts[1].strip() if len(parts) > 1 else ''
            fields.append((label, value))

    if not fields:
        return False

    # Create two-column table
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(fields):
        # Label cell
        cell_label = table.cell(i, 0)
        cell_label.text = ''
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

        # Value cell
        cell_value = table.cell(i, 1)
        cell_value.text = ''
        p = cell_value.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.0)

    logger.info(f"  ✓ Synopsis table: {len(fields)} fields")
    return True


def _add_abbreviation_table(doc: Document, usdm: Dict) -> bool:
    """Render abbreviations as a proper 2-column DOCX table for §13 Glossary.

    Returns True if the table was added, False if no abbreviation data.
    """
    study = usdm.get('study', {})
    versions = study.get('versions', [{}])
    version = versions[0] if versions else {}

    abbreviations = version.get('abbreviations', [])
    if not abbreviations:
        return False

    # Filter and sort
    valid = [a for a in abbreviations if isinstance(a, dict)]
    valid = sorted(
        valid,
        key=lambda a: (a.get('abbreviatedText', '') or a.get('name', '')).upper()
    )
    if not valid:
        return False

    # Create table: header row + data rows
    table = doc.add_table(rows=1 + len(valid), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, header_text in enumerate(['Abbreviation', 'Term']):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        # Grey shading
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F2F2F2',
        })
        tc_pr.append(shd)

    # Mark header as repeating
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    hdr_el = tr_pr.makeelement(qn('w:tblHeader'), {})
    tr_pr.append(hdr_el)

    # Data rows
    for ri, abbr in enumerate(valid, 1):
        short = abbr.get('abbreviatedText', abbr.get('name', ''))
        expanded = abbr.get('expansionText', abbr.get('text', ''))

        cell_abbr = table.rows[ri].cells[0]
        cell_abbr.text = ''
        p = cell_abbr.paragraphs[0]
        run = p.add_run(short or '')
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

        cell_term = table.rows[ri].cells[1]
        cell_term.text = ''
        p = cell_term.paragraphs[0]
        run = p.add_run(expanded or '')
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    logger.info(f"  ✓ Abbreviation table: {len(valid)} entries")
    return True
