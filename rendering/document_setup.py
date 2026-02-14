"""
M11 Document Setup — styles, TOC, headers/footers, title page.

Extracted from m11_renderer.py to reduce monolith size.
All functions operate on a python-docx Document instance.

Reference: ICH M11 Guideline, Template & Technical Specification (Step 4, 19 Nov 2025)
"""

import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _setup_styles(doc: Document) -> None:
    """Configure document styles per ICH M11 Template conventions.

    ICH M11 Template (Step 4, Nov 2025) heading table specifies:
      - Font: Times New Roman throughout
      - Body: 11pt, 1.15 line spacing, 6pt after paragraph
      - L1 (§N): 14pt TIMES NEW ROMAN BOLD BLACK (ALL CAPS)
      - L2 (§N.N): 14pt Times New Roman Bold Black
      - L3 (§N.N.N): 12pt Times New Roman Bold Black
      - L4 (§N.N.N.N): 12pt Times New Roman Bold Black
      - Non-numbered headings: 12pt Times New Roman Bold Black
      - Margins: 1 inch (2.54 cm) all sides
      - Page size: Letter (8.5 x 11 in)
    """
    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Normal (body text) ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # --- Heading 1 (M11 major sections: §1, §2, ... §14) ---
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.all_caps = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    # --- Heading 2 (M11 sub-sections: §1.1, §3.1, etc.) ---
    # M11 Template: 14 point Times New Roman Bold Black
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.font.all_caps = False
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # --- Heading 3 (M11 sub-sub-sections: §1.1.2, §5.1.1, etc.) ---
    # M11 Template: 12 point Times New Roman Bold Black
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # --- Heading 4 (§N.N.N.N — e.g. 10.4.1, 6.10.1) ---
    # M11 Template: 12 point Times New Roman Bold Black
    try:
        h4 = doc.styles['Heading 4']
    except KeyError:
        h4 = doc.styles.add_style('Heading 4', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    h4.font.name = 'Times New Roman'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.italic = False
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.space_before = Pt(6)
    h4.paragraph_format.space_after = Pt(3)
    h4.paragraph_format.keep_with_next = True

    # --- Heading 5 (§N.N.N.N.N — e.g. 10.4.1.1, 9.1.3.1) ---
    # M11 uses same 12pt bold for these deep sub-sections
    try:
        h5 = doc.styles['Heading 5']
    except KeyError:
        h5 = doc.styles.add_style('Heading 5', 1)
    h5.font.name = 'Times New Roman'
    h5.font.size = Pt(12)
    h5.font.bold = True
    h5.font.italic = False
    h5.font.color.rgb = RGBColor(0, 0, 0)
    h5.paragraph_format.space_before = Pt(6)
    h5.paragraph_format.space_after = Pt(3)
    h5.paragraph_format.keep_with_next = True

    # --- List Bullet ---
    try:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Times New Roman'
        lb.font.size = Pt(11)
        lb.paragraph_format.space_after = Pt(3)
        lb.paragraph_format.space_before = Pt(0)
        lb.paragraph_format.left_indent = Inches(0.5)
    except KeyError:
        pass

    # --- List Number ---
    try:
        ln = doc.styles['List Number']
        ln.font.name = 'Times New Roman'
        ln.font.size = Pt(11)
        ln.paragraph_format.space_after = Pt(3)
        ln.paragraph_format.space_before = Pt(0)
        ln.paragraph_format.left_indent = Inches(0.5)
    except KeyError:
        pass


def _add_toc_field(doc: Document) -> None:
    """Insert a real Word TOC field code.

    When the document is opened in Word, right-click the TOC and select
    "Update Field" to populate it with actual page numbers.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    # TOC field: { TOC \o "1-3" \h \z \u }
    run1 = p.add_run()
    fldChar_begin = run1._element.makeelement(
        qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}
    )
    run1._element.append(fldChar_begin)

    run2 = p.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar_separate = run3._element.makeelement(
        qn('w:fldChar'), {qn('w:fldCharType'): 'separate'}
    )
    run3._element.append(fldChar_separate)

    run4 = p.add_run('Right-click here and select "Update Field" to populate the Table of Contents')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar_end = run5._element.makeelement(
        qn('w:fldChar'), {qn('w:fldCharType'): 'end'}
    )
    run5._element.append(fldChar_end)


def _add_headers_footers(doc: Document, usdm: Dict) -> None:
    """Add M11-compliant headers and footers.

    Header: Sponsor Protocol Identifier + Confidential (right-aligned)
    Footer: Page N of M (centered), Version/Date (right-aligned)
    """
    study = usdm.get('study', {})
    versions = study.get('versions', [{}])
    version = versions[0] if versions else {}

    # Extract sponsor protocol ID
    identifiers = version.get('studyIdentifiers', study.get('studyIdentifiers', []))
    protocol_id = ''
    for ident in identifiers:
        if not isinstance(ident, dict):
            continue
        ident_type = ident.get('type', {})
        code = ident_type.get('code', '') if isinstance(ident_type, dict) else ''
        decode = ident_type.get('decode', '') if isinstance(ident_type, dict) else ''
        value = ident.get('text', ident.get('studyIdentifier', ''))
        if (code == 'C132351' or 'sponsor' in decode.lower()) and value:
            protocol_id = value
            break

    ver_text = version.get('versionIdentifier', '')
    date_text = version.get('effectiveDate', '')
    version_line = f'Version {ver_text}' if ver_text else ''
    if date_text:
        version_line += f'  |  {date_text}' if version_line else date_text

    # Apply headers/footers to all document sections (portrait + landscape)
    for section in doc.sections:
        section.different_first_page_header_footer = False

        # --- Header ---
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ''
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)

        # Left: Protocol ID
        if protocol_id:
            run_pid = hp.add_run(protocol_id)
            run_pid.font.size = Pt(8)
            run_pid.font.name = 'Times New Roman'

        # Right-aligned tab for Confidential
        hp.add_run('\t\t')
        run_conf = hp.add_run('CONFIDENTIAL')
        run_conf.bold = True
        run_conf.font.size = Pt(8)
        run_conf.font.name = 'Times New Roman'
        run_conf.font.color.rgb = RGBColor(192, 0, 0)

        # Add tab stops (center + right)
        pPr = hp._element.get_or_add_pPr()
        tabs = pPr.makeelement(qn('w:tabs'), {})
        # Center tab at 3.25"
        tab_center = tabs.makeelement(qn('w:tab'), {
            qn('w:val'): 'center',
            qn('w:pos'): str(int(Emu(Inches(3.25)))),
        })
        tabs.append(tab_center)
        # Right tab at 6.5"
        tab_right = tabs.makeelement(qn('w:tab'), {
            qn('w:val'): 'right',
            qn('w:pos'): str(int(Emu(Inches(6.5)))),
        })
        tabs.append(tab_right)
        pPr.append(tabs)

        # --- Footer ---
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ''
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)

        # Add tab stops for footer
        pPr2 = fp._element.get_or_add_pPr()
        tabs2 = pPr2.makeelement(qn('w:tabs'), {})
        tab_center2 = tabs2.makeelement(qn('w:tab'), {
            qn('w:val'): 'center',
            qn('w:pos'): str(int(Emu(Inches(3.25)))),
        })
        tabs2.append(tab_center2)
        tab_right2 = tabs2.makeelement(qn('w:tab'), {
            qn('w:val'): 'right',
            qn('w:pos'): str(int(Emu(Inches(6.5)))),
        })
        tabs2.append(tab_right2)
        pPr2.append(tabs2)

        # Left: version info
        if version_line:
            run_ver = fp.add_run(version_line)
            run_ver.font.size = Pt(8)
            run_ver.font.name = 'Times New Roman'
            run_ver.font.color.rgb = RGBColor(128, 128, 128)

        # Center: Page N of M (Word field codes)
        fp.add_run('\t')
        run_page_label = fp.add_run('Page ')
        run_page_label.font.size = Pt(8)
        run_page_label.font.name = 'Times New Roman'

        # PAGE field
        run_page_num = fp.add_run()
        run_page_num.font.size = Pt(8)
        fldBegin = run_page_num._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run_page_num._element.append(fldBegin)
        run_instr = fp.add_run()
        instrText = run_instr._element.makeelement(qn('w:instrText'), {})
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run_instr._element.append(instrText)
        run_sep = fp.add_run()
        fldSep = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
        run_sep._element.append(fldSep)
        run_num = fp.add_run('1')
        run_num.font.size = Pt(8)
        run_end = fp.add_run()
        fldEnd = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run_end._element.append(fldEnd)

        run_of = fp.add_run(' of ')
        run_of.font.size = Pt(8)
        run_of.font.name = 'Times New Roman'

        # NUMPAGES field
        run_total_num = fp.add_run()
        run_total_num.font.size = Pt(8)
        fldBegin2 = run_total_num._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run_total_num._element.append(fldBegin2)
        run_instr2 = fp.add_run()
        instrText2 = run_instr2._element.makeelement(qn('w:instrText'), {})
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = ' NUMPAGES '
        run_instr2._element.append(instrText2)
        run_sep2 = fp.add_run()
        fldSep2 = run_sep2._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
        run_sep2._element.append(fldSep2)
        run_total_num = fp.add_run('1')
        run_total_num.font.size = Pt(8)
        run_end2 = fp.add_run()
        fldEnd2 = run_end2._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run_end2._element.append(fldEnd2)


def _add_title_page(doc: Document, usdm: Dict) -> None:
    """Add an M11-compliant title page with all Required/Optional fields.

    M11 Title Page data elements (from Technical Specification):
      Required: Full Title, Sponsor Protocol Identifier, Original Protocol
                Indicator, Trial Phase, Sponsor Name and Address,
                Regulatory/Clinical Trial Identifiers, Sponsor Approval
      Optional: Confidentiality Statement, Trial Acronym, Short Title,
                Version Number/Date, Amendment details, IP codes/names,
                Co-Sponsor, Signatory, Medical Expert Contact
    """
    study = usdm.get('study', {})
    versions = study.get('versions', [{}])
    version = versions[0] if versions else {}
    titles = version.get('titles', [])
    study_design = (version.get('studyDesigns', [{}]) or [{}])[0]

    # ---- Extract title variants by type ----
    full_title = ''
    short_title = ''
    acronym = ''
    for t in titles:
        if not isinstance(t, dict):
            continue
        ttype = t.get('type', {})
        decode = ttype.get('decode', '') if isinstance(ttype, dict) else ''
        code = ttype.get('code', '') if isinstance(ttype, dict) else ''
        txt = t.get('text', '')
        if not txt:
            continue
        dl = decode.lower()
        if 'official' in dl or (not full_title and len(txt) > len(full_title)):
            full_title = txt
        if 'brief' in dl or 'short' in dl:
            short_title = txt
        if 'acronym' in dl:
            acronym = txt

    if not full_title:
        full_title = 'Clinical Protocol'

    # ---- Extract identifiers by C-code / decode ----
    # C-code mapping for M11 regulatory identifiers
    IDENT_MAP = {
        'C132351': 'Sponsor Protocol Identifier',
        'C172240': 'NCT Number',
        'C218684': 'EU CT Number',
        'C218685': 'FDA IND Number',
        'C218686': 'IDE Number',
        'C218687': 'jRCT Number',
        'C218688': 'NMPA IND Number',
        'C218689': 'WHO/UTN Number',
        'C98714':  'Registry Identifier',
    }

    identifiers = version.get('studyIdentifiers', study.get('studyIdentifiers', []))
    sponsor_protocol_id = ''
    regulatory_ids: List[tuple] = []

    for ident in identifiers:
        if not isinstance(ident, dict):
            continue
        ident_type = ident.get('type', {})
        code = ident_type.get('code', '') if isinstance(ident_type, dict) else ''
        decode = ident_type.get('decode', '') if isinstance(ident_type, dict) else ''
        value = ident.get('text', ident.get('studyIdentifier', ''))
        if not value:
            continue

        if code == 'C132351' or 'sponsor' in decode.lower():
            sponsor_protocol_id = value
        elif code in IDENT_MAP:
            regulatory_ids.append((IDENT_MAP[code], value))
        elif 'registry' in decode.lower() or 'identifier' in decode.lower():
            regulatory_ids.append((decode, value))

    # ---- Extract organizations ----
    organizations = version.get('organizations', study.get('organizations', []))
    sponsor_name = ''
    sponsor_address = ''
    for org in organizations:
        if not isinstance(org, dict):
            continue
        org_type = org.get('type', {})
        org_decode = org_type.get('decode', '') if isinstance(org_type, dict) else ''
        org_code = org_type.get('code', '') if isinstance(org_type, dict) else ''
        # C54149 = Pharmaceutical Company (USDM CT C188724), C70793 = Clinical Study Sponsor (legacy)
        if org_code in ('C54149', 'C70793') or 'sponsor' in org_decode.lower() or 'pharma' in org_decode.lower():
            sponsor_name = org.get('name', '')
            sponsor_address = org.get('legalAddress', org.get('address', ''))
            if isinstance(sponsor_address, dict):
                parts = []
                for k in ['line', 'city', 'district', 'state', 'postalCode', 'country']:
                    val = sponsor_address.get(k, '')
                    if not val:
                        continue
                    if isinstance(val, dict):
                        val = val.get('decode', val.get('code', ''))
                    if val:
                        parts.append(str(val))
                sponsor_address = ', '.join(parts) if parts else sponsor_address.get('text', '')
            break

    # ---- Trial Phase ----
    # studyPhase may be an AliasCode (standardCode.decode) or flat Code (decode)
    phase = version.get('studyPhase', study_design.get('studyPhase', {}))
    phase_text = ''
    if isinstance(phase, dict):
        std = phase.get('standardCode', {})
        if isinstance(std, dict) and std.get('decode'):
            phase_text = std['decode']
        else:
            phase_text = phase.get('decode', phase.get('code', ''))

    # ---- Investigational Product info ----
    interventions = version.get('studyInterventions', [])
    ip_names = []
    for si in interventions:
        if not isinstance(si, dict):
            continue
        role = si.get('role', si.get('type', {}))
        role_decode = role.get('decode', '') if isinstance(role, dict) else ''
        if 'investigational' in role_decode.lower() or not ip_names:
            name = si.get('name', '')
            if name:
                ip_names.append(name)

    # ---- Amendments ----
    amendments = version.get('amendments', [])
    has_amendment = bool(amendments)
    latest_amendment = amendments[-1] if amendments else {}

    # ---- Version info ----
    ver_text = version.get('versionIdentifier', '')
    date_text = version.get('effectiveDate', '')

    # ================================================================
    # Render title page
    # ================================================================

    # Confidentiality statement (M11 Optional) — top of page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('CONFIDENTIAL')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(192, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(
        'This document is confidential and the property of the sponsor. '
        'No part of it may be transmitted, reproduced, published, or used '
        'without prior written authorization from the sponsor.'
    )
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Vertical spacing before title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(0)

    # Full Title (M11 Required)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(full_title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    # Short Title / Acronym line
    subtitle_parts = []
    if short_title:
        subtitle_parts.append(short_title)
    if acronym:
        subtitle_parts.append(f'({acronym})')
    if subtitle_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(' '.join(subtitle_parts))
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.italic = True

    # ---- Metadata table (M11 structured fields) ----
    meta_items: List[tuple] = []

    if sponsor_protocol_id:
        meta_items.append(('Sponsor Protocol Identifier', sponsor_protocol_id))
    meta_items.append(('Original Protocol', 'No' if has_amendment else 'Yes'))
    if ver_text:
        meta_items.append(('Version Number', ver_text))
    if date_text:
        meta_items.append(('Version Date', date_text))
    if has_amendment and isinstance(latest_amendment, dict):
        amend_name = latest_amendment.get('name', latest_amendment.get('number', ''))
        if amend_name:
            meta_items.append(('Amendment Identifier', amend_name))
        scope = latest_amendment.get('scope', {})
        scope_text = scope.get('decode', '') if isinstance(scope, dict) else ''
        if scope_text:
            meta_items.append(('Amendment Scope', scope_text))
    if ip_names:
        meta_items.append(("Investigational Product(s)", ', '.join(ip_names)))
    if phase_text:
        meta_items.append(('Trial Phase', phase_text))
    if sponsor_name:
        sponsor_val = sponsor_name
        if sponsor_address:
            sponsor_val += f'\n{sponsor_address}'
        meta_items.append(('Sponsor Name and Address', sponsor_val))
    for label, value in regulatory_ids:
        meta_items.append((label, value))
    # ---- Sponsor Approval Date from dateValues ----
    sponsor_approval_date = ''
    for dv in version.get('dateValues', []):
        if not isinstance(dv, dict):
            continue
        dv_name = dv.get('name', '').lower()
        dv_type = dv.get('type', {})
        dv_decode = dv_type.get('decode', '').lower() if isinstance(dv_type, dict) else ''
        if 'sponsor approval' in dv_name or 'sponsor approval' in dv_decode:
            sponsor_approval_date = dv.get('dateValue', '')
            break
    if not sponsor_approval_date:
        # Fallback: use any protocol date
        for dv in version.get('dateValues', []):
            if isinstance(dv, dict) and dv.get('dateValue'):
                sponsor_approval_date = dv['dateValue']
                break
    meta_items.append(('Sponsor Approval Date',
                        sponsor_approval_date or date_text or '[Date of sponsor approval]'))

    # Render the metadata table — borderless with light shading on labels
    if meta_items:
        table = doc.add_table(rows=len(meta_items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Use no-border style for a cleaner look
        table.style = 'Table Grid'

        for i, (label, value) in enumerate(meta_items):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            cell_label.text = ''
            cell_value.text = ''

            # Label cell — bold, light grey background
            p_label = cell_label.paragraphs[0]
            p_label.paragraph_format.space_before = Pt(2)
            p_label.paragraph_format.space_after = Pt(2)
            run_label = p_label.add_run(label)
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_label.font.name = 'Times New Roman'
            # Grey background
            tcPr = cell_label._element.get_or_add_tcPr()
            shd = tcPr.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'})
            tcPr.append(shd)

            # Value cell
            p_value = cell_value.paragraphs[0]
            p_value.paragraph_format.space_before = Pt(2)
            p_value.paragraph_format.space_after = Pt(2)
            run_value = p_value.add_run(str(value))
            run_value.font.size = Pt(10)
            run_value.font.name = 'Times New Roman'

            # Set column widths
            cell_label.width = Cm(6)
            cell_value.width = Cm(10)

    # Generation timestamp (bottom of title page)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f'Document generated: {datetime.now().strftime("%Y-%m-%d %H:%M UTC")}')
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    run.font.color.rgb = RGBColor(160, 160, 160)

    doc.add_page_break()
