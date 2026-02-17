"""
M11 Text Formatting — narrative text rendering and content distribution.

Extracted from m11_renderer.py to reduce monolith size.
Handles markdown-style formatting, bullet/numbered lists, and
sub-section content distribution by keyword matching.
"""

import re
from typing import Dict, List

from docx import Document
from docx.shared import RGBColor


# XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
_ILLEGAL_XML_RE = re.compile(
    '[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f'
    '\ud800-\udfff\ufdd0-\ufdef\ufffe\uffff]'
)


def _sanitize_xml_text(text: str) -> str:
    """Strip characters that are illegal in XML 1.0 (used by python-docx)."""
    return _ILLEGAL_XML_RE.sub('', text)


def _add_narrative_text(doc: Document, text: str) -> None:
    """Add narrative text to the document with proper formatting.

    Handles:
      - Double-newline paragraph breaks
      - **bold** markers (markdown style)
      - Bullet lists (- item or • item)
      - Numbered lists (1. item, a. item)
      - Markdown headings (### Heading)
      - Regular paragraphs with preserved line breaks
    """
    text = _sanitize_xml_text(text)
    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        lines = para_text.split('\n')

        # Check if this block is a list (all lines start with - or •)
        is_bullet_list = all(
            l.strip().startswith('- ') or l.strip().startswith('• ')
            for l in lines if l.strip()
        )

        # Check if this block is a numbered list
        is_numbered_list = all(
            re.match(r'^\s*(\d+|[a-z])\.\s', l.strip())
            for l in lines if l.strip()
        )

        if is_bullet_list:
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Strip bullet prefix
                if line.startswith('- '):
                    line = line[2:]
                elif line.startswith('• '):
                    line = line[2:]
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_run(p, line)

        elif is_numbered_list:
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Strip number prefix but keep the text
                m = re.match(r'^\s*(?:\d+|[a-z])\.\s+(.*)', line)
                content = m.group(1) if m else line
                p = doc.add_paragraph(style='List Number')
                _add_formatted_run(p, content)

        elif para_text.startswith('### '):
            # Markdown L3 heading
            doc.add_heading(para_text[4:].strip(), level=3)

        elif para_text.startswith('## '):
            # Markdown L2 heading
            doc.add_heading(para_text[3:].strip(), level=2)

        elif '**' in para_text:
            # Contains bold markers — render with inline formatting
            p = doc.add_paragraph()
            _add_formatted_run(p, para_text)

        else:
            # Regular paragraph — preserve line breaks within
            p = doc.add_paragraph()
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                if i > 0:
                    p.add_run('\n')
                p.add_run(line)


def _add_formatted_run(paragraph, text: str) -> None:
    """Add text to a paragraph with inline **bold** and *italic* formatting."""
    text = _sanitize_xml_text(text)
    # Split on bold markers first
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            # Bold content — check for nested italic
            italic_parts = re.split(r'\*(.*?)\*', part)
            for j, ip in enumerate(italic_parts):
                if not ip:
                    continue
                run = paragraph.add_run(ip)
                run.bold = True
                if j % 2 == 1:
                    run.italic = True
        else:
            # Normal content — check for italic
            italic_parts = re.split(r'\*(.*?)\*', part)
            for j, ip in enumerate(italic_parts):
                if not ip:
                    continue
                run = paragraph.add_run(ip)
                if j % 2 == 1:
                    run.italic = True


def _distribute_to_subsections(
    narrative_text: str,
    subheadings: List[tuple],
) -> Dict[str, str]:
    """Distribute narrative paragraphs to M11 sub-sections by keyword matching.

    Splits the narrative into paragraphs and scores each against the sub-heading
    keyword lists.  Returns a dict mapping sub_number → matched text, plus
    '_general' for unmatched paragraphs.

    Args:
        narrative_text: Full narrative text for the section
        subheadings: List of (sub_number, title, level, keywords) tuples

    Returns:
        Dict mapping sub_number → concatenated paragraph text
    """
    paragraphs = [p.strip() for p in narrative_text.split('\n') if p.strip()]
    if not paragraphs:
        return {'_general': ''}

    # Build buckets
    buckets: Dict[str, List[str]] = {'_general': []}
    for sub_num, _title, _level, _kw in subheadings:
        buckets[sub_num] = []

    for para in paragraphs:
        para_lower = para.lower()
        best_sub = None
        best_score = 0

        for sub_num, _title, _level, keywords in subheadings:
            score = sum(1 for kw in keywords if kw.lower() in para_lower)
            if score > best_score:
                best_score = score
                best_sub = sub_num

        if best_sub and best_score > 0:
            buckets[best_sub].append(para)
        else:
            buckets['_general'].append(para)

    # Join each bucket
    return {k: '\n'.join(v) for k, v in buckets.items()}
