"""
M11 Protocol Document Renderer — generates DOCX from USDM JSON.

Transforms a protocol_usdm.json into an ICH M11-structured Word document
by combining:
  1. M11 section mapping (from m11_mapper)
  2. Narrative content (extracted text)
  3. Structured USDM entities (objectives, endpoints, arms, epochs, etc.)

The renderer produces a professional DOCX with:
  - Title page with protocol metadata
  - Table of contents placeholder
  - All 14 M11 sections with proper heading hierarchy
  - Auto-composed entity sections (objectives, eligibility, study design)
  - Synopsis table

Reference: ICH M11 Guideline, Template & Technical Specification (Step 4, 19 Nov 2025)
"""

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .document_setup import _setup_styles, _add_toc_field, _add_headers_footers, _add_title_page
from .text_formatting import _add_narrative_text, _add_formatted_run, _distribute_to_subsections

logger = logging.getLogger(__name__)


@dataclass
class M11RenderResult:
    """Result of M11 DOCX rendering."""
    success: bool
    output_path: Optional[str] = None
    error: Optional[str] = None
    sections_rendered: int = 0
    sections_with_content: int = 0
    total_words: int = 0


# Style helpers, TOC field, headers/footers, title page → rendering/document_setup.py
# Text formatting, subsection distribution → rendering/text_formatting.py


# ---------------------------------------------------------------------------
# Entity composers — imported from rendering.composers
# ---------------------------------------------------------------------------
from .composers import (
    _compose_synopsis,
    _compose_interventions,
    _compose_objectives,
    _compose_study_design,
    _compose_eligibility,
    _compose_estimands,
    _compose_statistics,
    _compose_safety,
    _compose_discontinuation,
    _compose_glossary,
    _compose_references,
)
from .tables import _add_soa_table, _add_synopsis_table, _add_abbreviation_table


# Legacy stubs removed — all composer functions now live in rendering/composers.py
# The following functions were moved:
#   _compose_synopsis        → §1.1.2 Overall Design
#   _compose_interventions   → §6 Trial Interventions
#   _compose_objectives      → §3 Objectives and Endpoints
#   _compose_study_design    → §4 Trial Design
#   _compose_eligibility     → §5 Study Population
#   _compose_estimands       → §3.1 Estimands
#   _compose_statistics      → §10 Statistical Considerations
#   _compose_safety          → §9 Adverse Events / Safety
#   _compose_discontinuation → §7 Discontinuation


def _find_schema_figure(usdm: Dict, output_dir: Optional[str]) -> Optional[str]:
    """Find a study/trial schema figure image path from USDM data.

    Searches protocolFigures (direct or extension attribute) for a figure
    labeled 'Study Schema', 'Trial Schema', 'Figure 1', or with
    sectionNumber '1.2' or '4'. Returns the absolute path to the rendered
    PNG if it exists on disk, else None.
    """
    if not output_dir:
        return None

    # Gather figures from multiple locations
    figures = usdm.get('protocolFigures', [])

    # Also check extension attributes
    if not figures:
        study = usdm.get('study', {})
        versions = study.get('versions', [{}])
        version = versions[0] if versions else {}
        for ext in version.get('extensionAttributes', []):
            if isinstance(ext, dict) and 'protocol-figures' in (ext.get('url', '') or ''):
                try:
                    figures = json.loads(ext.get('valueString', '[]'))
                except (json.JSONDecodeError, TypeError):
                    pass
                break

    if not figures:
        return None

    # Priority order for schema figure detection
    _SCHEMA_LABELS = {'study schema', 'trial schema', 'study design schema'}
    _SCHEMA_SECTIONS = {'1.2', '4', '4.1'}

    best = None
    for fig in figures:
        if not isinstance(fig, dict):
            continue
        label = (fig.get('label', '') or '').lower()
        title = (fig.get('title', '') or '').lower()
        sec = fig.get('sectionNumber', '') or ''
        img = fig.get('imagePath', '')

        if not img:
            continue

        # Exact match on label/title
        if label in _SCHEMA_LABELS or title in _SCHEMA_LABELS:
            best = img
            break
        if 'schema' in label or 'schema' in title:
            best = img
            break
        if sec in _SCHEMA_SECTIONS:
            best = img
            break
        # Fallback: Figure 1 is often the study schema
        if label == 'figure 1' and not best:
            best = img

    if not best:
        return None

    # Resolve to absolute path
    abs_path = Path(output_dir) / best
    if abs_path.exists():
        return str(abs_path)
    return None


# ---------------------------------------------------------------------------
# Main renderer
# ---------------------------------------------------------------------------

def render_m11_docx(
    usdm: Dict,
    output_path: str,
    m11_mapping: Optional[Dict] = None,
) -> M11RenderResult:
    """
    Render an ICH M11-structured DOCX from USDM JSON.

    Args:
        usdm: Full protocol_usdm.json dict
        output_path: Path for the output .docx file
        m11_mapping: Optional pre-computed M11 mapping from the pipeline.
                     If not provided, computes it from narrative data.

    Returns:
        M11RenderResult with success status and stats
    """
    result = M11RenderResult(success=False)
    output_dir = str(Path(output_path).parent) if output_path else None

    try:
        doc = Document()
        _setup_styles(doc)

        # Title page
        _add_title_page(doc, usdm)

        # Table of Contents — real Word TOC field
        doc.add_heading('Table of Contents', level=1)
        _add_toc_field(doc)
        doc.add_page_break()

        # Get or compute M11 mapping
        if m11_mapping is None:
            m11_mapping = _compute_m11_mapping(usdm)

        m11_sections = m11_mapping.get('sections', {})

        # Entity composers for enrichment
        entity_composers = {
            '1': _compose_synopsis,
            '3': _compose_objectives,
            '4': _compose_study_design,
            '5': _compose_eligibility,
            '6': _compose_interventions,
            '7': _compose_discontinuation,
            '9': _compose_safety,
            '10': _compose_statistics,
            '13': _compose_glossary,
            '14': _compose_references,
        }

        # Render each M11 section with sub-heading support
        from extraction.narrative.m11_mapper import M11_TEMPLATE, M11_SUBHEADINGS
        sections_rendered = 0
        sections_with_content = 0
        total_words = 0

        def _heading_level(section_number: str) -> int:
            """Determine Word heading level from M11 section number.

            M11 heading levels per template:
              '1'         → L1 (14pt ALL CAPS)
              '1.1'       → L2 (14pt bold)
              '1.1.2'     → L3 (12pt bold)
              '10.4.1'    → L4 (12pt bold)
              '10.4.1.1'  → L5 (12pt bold)
            Deeper levels capped at 5 (Word supports up to 9).
            """
            parts = section_number.split('.')
            return min(len(parts), 5)

        for m11 in M11_TEMPLATE:
            sec_data = m11_sections.get(m11.number, {})
            narrative_text = sec_data.get('text', '')
            has_content = bool(narrative_text.strip())

            # Add M11 section heading (L1 — style has all_caps=True)
            doc.add_heading(f'{m11.number} {m11.title}', level=1)

            # Distribute narrative to sub-sections if sub-headings defined
            subheadings = M11_SUBHEADINGS.get(m11.number, [])
            if narrative_text.strip() and subheadings:
                subsection_content = _distribute_to_subsections(
                    narrative_text, subheadings
                )
                # Render general content (not matched to any sub-heading)
                general = subsection_content.get('_general', '')
                if general.strip():
                    _add_narrative_text(doc, general)
                    total_words += len(general.split())
                # Render each sub-heading with its matched content
                for sub_num, sub_title, sub_level, _kw in subheadings:
                    sub_text = subsection_content.get(sub_num, '')
                    if sub_text.strip():
                        # Use section number depth for heading level
                        level = _heading_level(sub_num)
                        doc.add_heading(f'{sub_num} {sub_title}', level=level)
                        _add_narrative_text(doc, sub_text)
                        total_words += len(sub_text.split())
                    else:
                        # Render empty sub-section heading with placeholder
                        level = _heading_level(sub_num)
                        doc.add_heading(f'{sub_num} {sub_title}', level=level)
                sections_with_content += 1
            elif narrative_text.strip():
                # No sub-headings defined — render as flat narrative
                _add_narrative_text(doc, narrative_text)
                total_words += len(narrative_text.split())
                sections_with_content += 1

            # --- Section 1 special handling: Synopsis table + SoA table ---
            if m11.number == '1':
                # §1.1 Protocol Synopsis
                doc.add_heading('1.1 Protocol Synopsis', level=2)

                # §1.1.1 Primary and Secondary Objectives and Estimands
                doc.add_heading(
                    '1.1.1 Primary and Secondary Objectives and Estimands',
                    level=3
                )
                obj_text = _compose_objectives(usdm)
                if obj_text.strip():
                    _add_narrative_text(doc, obj_text)
                    total_words += len(obj_text.split())

                # §1.1.2 Overall Design — render as structured table
                doc.add_heading('1.1.2 Overall Design', level=3)
                synopsis_rendered = _add_synopsis_table(doc, usdm)
                if synopsis_rendered:
                    sections_with_content += 1
                elif not has_content:
                    composed = _compose_synopsis(usdm)
                    if composed.strip():
                        _add_narrative_text(doc, composed)
                        total_words += len(composed.split())
                        sections_with_content += 1

                # §1.2 Trial Schema — embed figure image if available
                doc.add_heading('1.2 Trial Schema', level=2)
                schema_fig = _find_schema_figure(usdm, output_dir)
                if schema_fig:
                    try:
                        doc.add_picture(schema_fig, width=Inches(6.0))
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        logger.warning(f"Failed to embed trial schema: {e}")
                        p = doc.add_paragraph()
                        run = p.add_run('[Trial schema diagram — image could not be embedded]')
                        run.italic = True
                        run.font.color.rgb = RGBColor(192, 0, 0)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run('[Trial schema diagram — refer to Section 4 '
                                    'Trial Design for details]')
                    run.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)

                # §1.3 Schedule of Activities — render as DOCX table
                doc.add_heading('1.3 Schedule of Activities', level=2)
                soa_rendered = _add_soa_table(doc, usdm)
                if soa_rendered:
                    if not has_content and not synopsis_rendered:
                        sections_with_content += 1
                else:
                    p = doc.add_paragraph()
                    run = p.add_run('[Schedule of Activities table not available — '
                                    'no schedule timeline data in USDM]')
                    run.italic = True
                    run.font.color.rgb = RGBColor(192, 0, 0)

            else:
                # Entity-composed content (supplement or replace empty narrative)
                composer = entity_composers.get(m11.number)
                composed = ''
                if composer:
                    composed = composer(usdm)
                    if composed.strip():
                        _add_narrative_text(doc, composed)
                        total_words += len(composed.split())
                        if not has_content:
                            sections_with_content += 1

                # Estimand tables for section 3
                if m11.number == '3':
                    estimand_text = _compose_estimands(usdm)
                    if estimand_text:
                        doc.add_heading(
                            '3.1 Primary Objective(s) and Associated Estimand(s)',
                            level=2
                        )
                        _add_narrative_text(doc, estimand_text)

                # Abbreviation table for section 13 (Glossary)
                abbr_rendered = False
                if m11.number == '13':
                    abbr_rendered = _add_abbreviation_table(doc, usdm)
                    if abbr_rendered and not has_content:
                        sections_with_content += 1

                # Empty section notice (reuse cached composed result)
                if not narrative_text.strip() and not composed.strip() and not abbr_rendered:
                    p = doc.add_paragraph()
                    run = p.add_run('[Section content not yet available]')
                    run.italic = True
                    run.font.color.rgb = RGBColor(192, 0, 0)

            sections_rendered += 1

            # Page break between major L1 sections (§1–§11)
            # Appendices §12–§14 flow continuously
            try:
                sec_num = int(m11.number)
                if sec_num <= 11:
                    doc.add_page_break()
            except ValueError:
                pass

        # Add headers and footers to all sections (must be after all
        # sections are created, including landscape SoA section)
        _add_headers_footers(doc, usdm)

        # Save
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        result.success = True
        result.output_path = output_path
        result.sections_rendered = sections_rendered
        result.sections_with_content = sections_with_content
        result.total_words = total_words

        logger.info(
            f"M11 DOCX rendered: {output_path} "
            f"({sections_with_content}/{sections_rendered} sections with content, "
            f"{total_words} words)"
        )

        # Run M11 conformance validation and save report
        try:
            from validation.m11_conformance import (
                validate_m11_conformance, save_conformance_report,
            )
            conf_report = validate_m11_conformance(usdm, m11_mapping)
            conf_path = str(Path(output_path).parent / 'm11_conformance_report.json')
            save_conformance_report(conf_report, conf_path)
        except Exception as conf_err:
            logger.warning(f"M11 conformance check failed: {conf_err}")

    except Exception as e:
        logger.error(f"M11 DOCX rendering failed: {e}")
        result.error = str(e)

    return result


def _compute_m11_mapping(usdm: Dict) -> Dict:
    """Compute M11 mapping from USDM narrative data on the fly."""
    from extraction.narrative.m11_mapper import map_sections_to_m11, build_m11_narrative

    study = usdm.get('study', {})
    versions = study.get('versions', [{}])
    version = versions[0] if versions else {}

    # Combine narrativeContents and narrativeContentItems
    narrative_contents = version.get('narrativeContents', [])
    narrative_items = version.get('narrativeContentItems', [])
    all_items = narrative_contents + narrative_items

    # Build section dicts with sectionType extraction
    sec_dicts = []
    sec_texts = {}
    seen_nums = set()
    for nc in all_items:
        if not isinstance(nc, dict):
            continue
        num = nc.get('sectionNumber', '')
        if not num or num in seen_nums:
            continue
        seen_nums.add(num)
        title = nc.get('sectionTitle', nc.get('name', ''))
        text = nc.get('text', '')
        # Extract section type for Pass 3
        sec_type = ''
        st = nc.get('sectionType', {})
        if isinstance(st, dict):
            sec_type = st.get('decode', st.get('code', ''))
        sec_dicts.append({
            'number': num,
            'title': title,
            'type': sec_type,
        })
        if text and text != title:
            sec_texts[num] = text

    if not sec_dicts:
        return {'sections': {}}

    mapping = map_sections_to_m11(sec_dicts, section_texts=sec_texts)
    m11_narrative = build_m11_narrative(sec_dicts, sec_texts, mapping)

    return {
        'coverage': f"{mapping.m11_covered}/{mapping.m11_total}",
        'requiredCoverage': f"{mapping.m11_required_covered}/{mapping.m11_required_total}",
        'sections': m11_narrative,
    }
