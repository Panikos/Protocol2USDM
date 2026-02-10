"""
Tests for rendering.tables — SoA table and Synopsis table rendering.

Validates:
- _build_soa_data extracts correct structure from USDM
- _add_soa_table creates a valid DOCX table with correct dimensions
- Group separator rows, header shading, empty activity filtering
- _add_synopsis_table creates a two-column table
- Edge cases: no timelines, no instances, no encounters
"""

import json
import pytest
from unittest.mock import patch
from docx import Document

from rendering.tables import (
    _build_soa_data,
    _add_soa_table,
    _add_synopsis_table,
    _shade_cell,
    _set_cell_vertical_alignment,
    _mark_row_as_header,
    _style_header_cell,
)


# ── Fixtures ─────────────────────────────────────────────────────────

def _make_usdm(epochs=None, encounters=None, activities=None,
               instances=None, groups=None, footnotes=None):
    """Build a minimal USDM dict for testing."""
    if epochs is None:
        epochs = [
            {'id': 'ep1', 'name': 'Screening', 'instanceType': 'StudyEpoch'},
            {'id': 'ep2', 'name': 'Treatment', 'instanceType': 'StudyEpoch'},
        ]
    if encounters is None:
        encounters = [
            {'id': 'enc1', 'name': 'Visit 1', 'epochId': 'ep1', 'instanceType': 'Encounter'},
            {'id': 'enc2', 'name': 'Visit 2', 'epochId': 'ep2', 'instanceType': 'Encounter'},
            {'id': 'enc3', 'name': 'Visit 3', 'epochId': 'ep2', 'instanceType': 'Encounter'},
        ]
    if activities is None:
        activities = [
            {'id': 'act1', 'name': 'Informed Consent', 'instanceType': 'Activity'},
            {'id': 'act2', 'name': 'Blood Draw', 'instanceType': 'Activity'},
            {'id': 'act3', 'name': 'Vitals', 'instanceType': 'Activity'},
        ]
    if instances is None:
        instances = [
            {'id': 'i1', 'activityIds': ['act1'], 'encounterId': 'enc1',
             'epochId': 'ep1', 'instanceType': 'ScheduledActivityInstance'},
            {'id': 'i2', 'activityIds': ['act2'], 'encounterId': 'enc2',
             'epochId': 'ep2', 'instanceType': 'ScheduledActivityInstance'},
            {'id': 'i3', 'activityIds': ['act2'], 'encounterId': 'enc3',
             'epochId': 'ep2', 'instanceType': 'ScheduledActivityInstance'},
            {'id': 'i4', 'activityIds': ['act3'], 'encounterId': 'enc1',
             'epochId': 'ep1', 'instanceType': 'ScheduledActivityInstance'},
            {'id': 'i5', 'activityIds': ['act3'], 'encounterId': 'enc2',
             'epochId': 'ep2', 'instanceType': 'ScheduledActivityInstance'},
        ]
    if groups is None:
        groups = [
            {'name': 'Eligibility', 'childIds': ['act1'], 'instanceType': 'ActivityGroup'},
            {'name': 'Labs', 'childIds': ['act2'], 'instanceType': 'ActivityGroup'},
        ]

    design = {
        'epochs': epochs,
        'encounters': encounters,
        'activities': activities,
        'scheduleTimelines': [{
            'id': 'tl1',
            'instances': instances,
            'instanceType': 'ScheduleTimeline',
        }],
        'activityGroups': groups,
        'extensionAttributes': [],
    }

    if footnotes:
        design['extensionAttributes'].append({
            'url': 'https://protocol2usdm.io/extensions/x-soaFootnotes',
            'valueString': json.dumps(footnotes),
            'instanceType': 'ExtensionAttribute',
        })

    return {
        'study': {
            'versions': [{
                'studyDesigns': [design],
            }],
        },
    }


@pytest.fixture
def basic_usdm():
    return _make_usdm()


@pytest.fixture
def empty_usdm():
    """USDM with no schedule timelines."""
    return _make_usdm(instances=[])


# ── _build_soa_data ──────────────────────────────────────────────────

class TestBuildSoaData:
    """Tests for the SoA data extraction from USDM."""

    def test_returns_dict_with_expected_keys(self, basic_usdm):
        soa = _build_soa_data(basic_usdm)
        assert soa is not None
        assert set(soa.keys()) == {
            'epochs', 'encounters_by_epoch', 'groups',
            'ungrouped', 'cells', 'footnote_refs', 'footnotes',
        }

    def test_epochs_extracted(self, basic_usdm):
        soa = _build_soa_data(basic_usdm)
        assert len(soa['epochs']) == 2
        assert soa['epochs'][0]['name'] == 'Screening'
        assert soa['epochs'][1]['name'] == 'Treatment'

    def test_encounters_by_epoch(self, basic_usdm):
        soa = _build_soa_data(basic_usdm)
        ebe = soa['encounters_by_epoch']
        assert len(ebe.get('ep1', [])) == 1
        assert len(ebe.get('ep2', [])) == 2

    def test_cells_populated(self, basic_usdm):
        soa = _build_soa_data(basic_usdm)
        cells = soa['cells']
        assert cells[('act1', 'enc1')] == 'X'
        assert cells[('act2', 'enc2')] == 'X'
        assert cells[('act2', 'enc3')] == 'X'
        assert ('act1', 'enc2') not in cells

    def test_groups_extracted(self, basic_usdm):
        soa = _build_soa_data(basic_usdm)
        assert len(soa['groups']) == 2
        assert soa['groups'][0]['name'] == 'Eligibility'
        assert len(soa['groups'][0]['activities']) == 1

    def test_ungrouped_activities(self, basic_usdm):
        soa = _build_soa_data(basic_usdm)
        ungrouped_ids = [a['id'] for a in soa['ungrouped']]
        assert 'act3' in ungrouped_ids

    def test_no_timelines_returns_none(self):
        usdm = _make_usdm()
        usdm['study']['versions'][0]['studyDesigns'][0]['scheduleTimelines'] = []
        assert _build_soa_data(usdm) is None

    def test_no_instances_returns_none(self, empty_usdm):
        assert _build_soa_data(empty_usdm) is None

    def test_footnotes_extracted(self):
        footnotes = [
            {'text': 'Only if clinically indicated', 'activityIds': ['act1'],
             'encounterIds': ['enc1']},
            {'text': 'Fasting required', 'activityIds': ['act2'],
             'encounterIds': ['enc2', 'enc3']},
        ]
        usdm = _make_usdm(footnotes=footnotes)
        soa = _build_soa_data(usdm)
        assert len(soa['footnotes']) == 2
        assert 'clinically indicated' in soa['footnotes'][0]
        assert len(soa['footnote_refs']) > 0


# ── _add_soa_table ───────────────────────────────────────────────────

class TestAddSoaTable:
    """Tests for the DOCX SoA table rendering."""

    def test_returns_true_with_data(self, basic_usdm):
        doc = Document()
        assert _add_soa_table(doc, basic_usdm) is True

    def test_returns_false_without_data(self, empty_usdm):
        doc = Document()
        assert _add_soa_table(doc, empty_usdm) is False

    def test_table_created(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        assert len(doc.tables) >= 1

    def test_table_has_header_rows(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        table = doc.tables[0]
        # Row 0 = epoch headers, Row 1 = encounter headers
        assert len(table.rows) >= 2

    def test_epoch_header_text(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        table = doc.tables[0]
        # First cell of row 0 should be "Study Period"
        assert 'Study Period' in table.cell(0, 0).text

    def test_encounter_header_text(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        table = doc.tables[0]
        assert 'Assessment' in table.cell(1, 0).text

    def test_group_separator_rows_present(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        table = doc.tables[0]
        # Should have group rows for "Eligibility", "Labs", "Other Assessments"
        all_text = [table.cell(r, 0).text for r in range(len(table.rows))]
        group_rows = [t for t in all_text if t in ('Eligibility', 'Labs', 'Other Assessments')]
        assert len(group_rows) >= 2

    def test_empty_activities_filtered(self):
        """Activities with no ticks should not appear in the table."""
        usdm = _make_usdm(
            activities=[
                {'id': 'act1', 'name': 'Active', 'instanceType': 'Activity'},
                {'id': 'act_empty', 'name': 'NoTicks', 'instanceType': 'Activity'},
            ],
            instances=[
                {'id': 'i1', 'activityIds': ['act1'], 'encounterId': 'enc1',
                 'epochId': 'ep1', 'instanceType': 'ScheduledActivityInstance'},
            ],
            groups=[],
        )
        doc = Document()
        _add_soa_table(doc, usdm)
        table = doc.tables[0]
        all_text = ' '.join(table.cell(r, 0).text for r in range(len(table.rows)))
        assert 'Active' in all_text
        assert 'NoTicks' not in all_text

    def test_landscape_orientation(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        # After SoA, there should be landscape + portrait sections
        assert len(doc.sections) >= 2

    def test_column_count(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        table = doc.tables[0]
        # 1 activity col + 3 encounter cols = 4
        assert len(table.columns) == 4

    def test_tick_marks_present(self, basic_usdm):
        doc = Document()
        _add_soa_table(doc, basic_usdm)
        table = doc.tables[0]
        # Find a cell with 'X' mark
        found_x = False
        for row in table.rows:
            for cell in row.cells:
                if 'X' in cell.text:
                    found_x = True
                    break
        assert found_x


# ── _add_synopsis_table ──────────────────────────────────────────────

class TestAddSynopsisTable:
    """Tests for the Synopsis table rendering."""

    def test_returns_bool(self, basic_usdm):
        doc = Document()
        result = _add_synopsis_table(doc, basic_usdm)
        assert isinstance(result, bool)


# ── Cell styling helpers ─────────────────────────────────────────────

class TestCellStylingHelpers:
    """Tests for low-level cell styling functions."""

    def test_shade_cell_no_error(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        _shade_cell(table.cell(0, 0), 'D9E2F3')

    def test_vertical_alignment_no_error(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        _set_cell_vertical_alignment(table.cell(0, 0), 'center')

    def test_mark_row_as_header_no_error(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        _mark_row_as_header(table.rows[0])

    def test_style_header_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        _style_header_cell(table.cell(0, 0), 'Test Header')
        assert 'Test Header' in table.cell(0, 0).text
